"""Module 1 report generator — DOCX and PDF output.

Produces a structured diagnostic report containing:
  - Title page with territory name, date and key indicators
  - Section 1: PA network map + coverage tables
  - Section 2: Ecosystem representativity chart + RI table
  - Section 3: Gap analysis map + gap metrics
  - Section 4: Criterion profiles chart + summary table

Static map images are generated with matplotlib/geopandas (folium maps
are interactive-only and cannot be embedded directly in documents).
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

import matplotlib
matplotlib.use('Agg')   # non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import pandas as pd
import geopandas as gpd
from shapely.geometry import box

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Colour palette (mirrors iucn_classification.yaml defaults)
# ---------------------------------------------------------------------------
_IUCN_COLOURS = {
    'strict_core':   '#2E7D32',
    'regulatory':    '#66BB6A',
    'contractual':   '#A5D6A7',
    'unassigned':    '#B4B2A9',
    'strict_gaps':   '#D0021B',
    'qualitative_gaps': '#F6A623',
    'corridors':     '#4A90E2',
}


# ---------------------------------------------------------------------------
# Static map helpers
# ---------------------------------------------------------------------------

def _pa_map_figure(
    pa_gdf: gpd.GeoDataFrame,
    territory_geom,
    iucn_classes: dict,
    figsize: tuple = (12, 8),
) -> plt.Figure:
    """Render PA network as a static matplotlib figure."""
    fig, ax = plt.subplots(figsize=figsize)

    # Territory boundary
    territory_gs = gpd.GeoSeries([territory_geom], crs='EPSG:3035')
    territory_gs.boundary.plot(ax=ax, color='black', linewidth=1.5, linestyle='--', label='Territory')

    # PA polygons coloured by protection_class
    if 'protection_class' in pa_gdf.columns:
        for pclass, group in pa_gdf.groupby('protection_class'):
            colour = iucn_classes.get(pclass, {}).get('colour', '#B4B2A9')
            label  = iucn_classes.get(pclass, {}).get('label', pclass)
            group.plot(ax=ax, color=colour, alpha=0.6, label=label)
    else:
        pa_gdf.plot(ax=ax, color='#66BB6A', alpha=0.6, label='Protected areas')

    ax.set_axis_off()
    ax.set_title('Protected Area Network', fontsize=14, fontweight='bold', pad=10)
    ax.legend(loc='lower right', fontsize=9, framealpha=0.9)
    fig.tight_layout()
    return fig


def _gap_map_figure(
    gap_layers: dict,
    territory_geom,
    figsize: tuple = (12, 8),
) -> plt.Figure:
    """Render gap layers as a static matplotlib figure."""
    fig, ax = plt.subplots(figsize=figsize)

    territory_gs = gpd.GeoSeries([territory_geom], crs='EPSG:3035')
    territory_gs.boundary.plot(ax=ax, color='black', linewidth=1.5, linestyle='--', label='Territory')

    _layer_style = [
        ('strict_gaps',       '#D0021B', 'Strict Gaps'),
        ('qualitative_gaps',  '#F6A623', 'Qualitative Gaps'),
        ('corridors',         '#4A90E2', 'Potential Corridors'),
    ]
    for key, colour, label in _layer_style:
        gdf = gap_layers.get(key)
        if gdf is not None and len(gdf) > 0:
            clean = gdf[~gdf.geometry.is_empty & gdf.geometry.notnull()]
            if len(clean) > 0:
                clean.plot(ax=ax, color=colour, alpha=0.5, label=label)

    ax.set_axis_off()
    ax.set_title('Gap Analysis', fontsize=14, fontweight='bold', pad=10)
    ax.legend(loc='lower right', fontsize=9, framealpha=0.9)
    fig.tight_layout()
    return fig


def _criterion_bar_figure(zonal_df: pd.DataFrame, figsize: tuple = (12, 6)) -> plt.Figure:
    """Render criterion scores by IUCN category as a grouped bar chart."""
    chart_data = zonal_df[zonal_df['criterion'] != 'landuse'][
        ['criterion', 'iucn_cat', 'mean']
    ].copy()

    # Normalise per criterion
    for crit in chart_data['criterion'].unique():
        mask = chart_data['criterion'] == crit
        vals = chart_data.loc[mask, 'mean']
        vmin, vmax = vals.min(), vals.max()
        if vmax > vmin:
            chart_data.loc[mask, 'mean'] = (vals - vmin) / (vmax - vmin)
        if crit == 'anthropogenic_pressure':
            chart_data.loc[mask, 'mean'] = 1.0 - chart_data.loc[mask, 'mean']

    criteria   = sorted(chart_data['criterion'].unique())
    iucn_cats  = [c for c in sorted(chart_data['iucn_cat'].unique()) if c != 'outside']
    if 'outside' in chart_data['iucn_cat'].unique():
        iucn_cats.append('outside')

    n_criteria = len(criteria)
    n_cats     = len(iucn_cats)
    x          = np.arange(n_criteria)
    width      = 0.8 / max(n_cats, 1)
    colours    = plt.cm.Set2(np.linspace(0, 1, n_cats))

    fig, ax = plt.subplots(figsize=figsize)
    for i, cat in enumerate(iucn_cats):
        vals = [
            chart_data.loc[
                (chart_data['criterion'] == c) & (chart_data['iucn_cat'] == cat), 'mean'
            ].values[0] if len(chart_data.loc[
                (chart_data['criterion'] == c) & (chart_data['iucn_cat'] == cat)
            ]) > 0 else 0.0
            for c in criteria
        ]
        ax.bar(x + i * width, vals, width, label=cat, color=colours[i], alpha=0.85)

    ax.set_xticks(x + width * (n_cats - 1) / 2)
    ax.set_xticklabels([c.replace('_', ' ').title() for c in criteria], rotation=30, ha='right')
    ax.set_ylim(0, 1.05)
    ax.set_ylabel('Normalised Score [0–1]')
    ax.set_title('Mean Criterion Scores by IUCN Category', fontweight='bold')
    ax.legend(title='IUCN Category', loc='upper right', fontsize=9)
    ax.axhline(0.5, color='grey', linewidth=0.8, linestyle='--', alpha=0.6)
    fig.tight_layout()
    return fig


def _ri_bar_figure(ri_df: pd.DataFrame, figsize: tuple = (10, 5)) -> plt.Figure:
    """Render ecosystem representativity as a horizontal bar chart."""
    df = ri_df[['ecosystem_type', 'coverage_pct']].sort_values('coverage_pct', ascending=True)
    colours = ['#1D9E75' if v >= 30 else '#F6A623' for v in df['coverage_pct']]

    fig, ax = plt.subplots(figsize=figsize)
    bars = ax.barh(df['ecosystem_type'], df['coverage_pct'], color=colours, alpha=0.85)
    ax.axvline(30, color='red', linewidth=1.2, linestyle='--', label='30% KMGBF target')
    ax.set_xlabel('Coverage (%)')
    ax.set_title('Ecosystem Representativity', fontweight='bold')
    ax.legend(fontsize=9)
    fig.tight_layout()
    return fig


def _fig_to_bytes(fig: plt.Figure, dpi: int = 150) -> bytes:
    """Render matplotlib figure to PNG bytes."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# DOCX report
# ---------------------------------------------------------------------------

def generate_docx_report(
    territory_name: str,
    territory_area_ha: float,
    pa_gdf: gpd.GeoDataFrame,
    territory_geom,
    iucn_classes: dict,
    coverage_df: pd.DataFrame,
    iucn_coverage_df: Optional[pd.DataFrame],
    gap_layers: Optional[dict],
    gap_stats: Optional[dict],
    ri_df: Optional[pd.DataFrame],
    zonal_df: Optional[pd.DataFrame],
    kmgbf_pct: float,
    net_area_ha: float,
    strict_pct: Optional[float] = None,
) -> bytes:
    """
    Generate a DOCX diagnostic report for Module 1.

    Returns
    -------
    bytes
        Raw bytes of the .docx file ready for st.download_button.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install with: pip install python-docx"
        ) from e

    doc = Document()

    # -----------------------------------------------------------------------
    # Document style helpers
    # -----------------------------------------------------------------------
    def _heading(text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def _add_table(df: pd.DataFrame, include_index: bool = False):
        """Add a pandas DataFrame as a styled Word table."""
        if include_index:
            df = df.reset_index()
        cols = list(df.columns)
        t = doc.add_table(rows=1 + len(df), cols=len(cols))
        t.style = 'Light List Accent 3'
        # Header row
        for j, col in enumerate(cols):
            cell = t.rows[0].cells[j]
            cell.text = str(col)
            run = cell.paragraphs[0].runs[0]
            run.bold = True
        # Data rows
        for i, row in df.iterrows():
            tr = t.rows[i + 1]
            for j, val in enumerate(row):
                tr.cells[j].text = str(val) if val is not None else ''
        return t

    def _add_image_bytes(img_bytes: bytes, width_inches: float = 6.0):
        buf = io.BytesIO(img_bytes)
        doc.add_picture(buf, width=Inches(width_inches))

    def _kv_table(rows: list[tuple[str, str]]):
        """Add a two-column key-value summary table."""
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = 'Light Shading'
        for i, (k, v) in enumerate(rows):
            t.rows[i].cells[0].text = k
            t.rows[i].cells[1].text = v
            t.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    # -----------------------------------------------------------------------
    # Title page
    # -----------------------------------------------------------------------
    title = doc.add_heading('Module 1 — Protection Network Diagnostic', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f'Territory: {territory_name}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)

    date_p = doc.add_paragraph(f'Generated: {datetime.now().strftime("%d %B %Y at %H:%M")}')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # Executive summary box
    _heading('Executive Summary', level=1)
    kv_rows = [
        ('Territory',                  territory_name),
        ('Territory area',             f'{territory_area_ha:,.0f} ha'),
        ('Net protected area',         f'{net_area_ha:,.0f} ha'),
        ('% protected — IUCN I–VI + OECMs (KMGBF T3)', f'{kmgbf_pct:.1f}%'),
        ('KMGBF Target 3 (30% by 2030)', '✓ Met' if kmgbf_pct >= 30 else '✗ Not yet met'),
    ]
    if strict_pct is not None:
        kv_rows.append(('  of which strict core only (IUCN I–II)', f'{strict_pct:.1f}%'))
    kv_rows += [
        ('Number of PA sites',  f'{len(pa_gdf):,}'),
        ('Report generated',    datetime.now().strftime('%Y-%m-%d %H:%M')),
    ]
    _kv_table(kv_rows)
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Section 1: PA Network
    # -----------------------------------------------------------------------
    doc.add_page_break()
    _heading('1. Protected Area Network', level=1)

    # Static map
    _heading('1.1 Network Map', level=2)
    try:
        fig_pa = _pa_map_figure(pa_gdf, territory_geom, iucn_classes)
        _add_image_bytes(_fig_to_bytes(fig_pa), width_inches=6.0)
    except Exception as e:
        doc.add_paragraph(f'[Map unavailable: {e}]')

    # Coverage by protection class
    _heading('1.2 Coverage by Protection Class', level=2)
    try:
        disp = coverage_df.copy()
        disp['area_ha']       = disp['area_ha'].apply(lambda x: f'{x:,.0f}')
        disp['pct_territory'] = disp['pct_territory'].apply(lambda x: f'{x:.2f}%')
        _add_table(disp)
    except Exception as e:
        doc.add_paragraph(f'[Table unavailable: {e}]')

    doc.add_paragraph()

    # Coverage by IUCN category
    if iucn_coverage_df is not None and len(iucn_coverage_df) > 0:
        _heading('1.3 Coverage by IUCN Category', level=2)
        try:
            _add_table(iucn_coverage_df)
        except Exception as e:
            doc.add_paragraph(f'[Table unavailable: {e}]')
        doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Section 2: Ecosystem Representativity
    # -----------------------------------------------------------------------
    if ri_df is not None and len(ri_df) > 0:
        doc.add_page_break()
        _heading('2. Ecosystem Representativity', level=1)

        doc.add_paragraph(
            'The Representativity Index (RI) measures the fraction of each ecosystem '
            'type that falls within the protected area network. RI = 1.0 means the '
            '30% KMGBF target is fully met for that ecosystem type.'
        )

        _heading('2.1 Representativity Chart', level=2)
        try:
            fig_ri = _ri_bar_figure(ri_df)
            _add_image_bytes(_fig_to_bytes(fig_ri), width_inches=6.0)
        except Exception as e:
            doc.add_paragraph(f'[Chart unavailable: {e}]')

        _heading('2.2 Representativity Index by Ecosystem Type', level=2)
        try:
            ri_disp = ri_df[['ecosystem_type', 'coverage_pct', 'RI', 'gap_ha']].copy()
            ri_disp['coverage_pct'] = ri_disp['coverage_pct'].apply(lambda x: f'{x:.1f}%')
            ri_disp['RI']           = ri_disp['RI'].apply(lambda x: f'{x:.3f}')
            ri_disp['gap_ha']       = ri_disp['gap_ha'].apply(lambda x: f'{x:,.0f}')
            ri_disp.columns         = ['Ecosystem Type', 'Coverage', 'RI', 'Gap (ha)']
            _add_table(ri_disp)
        except Exception as e:
            doc.add_paragraph(f'[Table unavailable: {e}]')
        doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Section 3: Gap Analysis
    # -----------------------------------------------------------------------
    if gap_layers is not None and gap_stats is not None:
        doc.add_page_break()
        _heading('3. Gap Analysis', level=1)

        def _pct(area_ha):
            return area_ha / territory_area_ha * 100.0 if territory_area_ha > 0 else 0.0

        _heading('3.1 Gap Metrics', level=2)
        _kv_table([
            ('Strict gaps',
             f"{gap_stats.get('strict_area', 0):,.0f} ha  "
             f"({_pct(gap_stats.get('strict_area', 0)):.1f}% of territory)"),
            ('Qualitative gaps',
             f"{gap_stats.get('qual_area', 0):,.0f} ha  "
             f"({_pct(gap_stats.get('qual_area', 0)):.1f}% of territory)"),
            ('Potential corridors',
             f"{gap_stats.get('corridor_area', 0):,.0f} ha  "
             f"({_pct(gap_stats.get('corridor_area', 0)):.1f}% of territory)"),
        ])
        doc.add_paragraph()

        _heading('3.2 Gap Layers Map', level=2)
        try:
            fig_gap = _gap_map_figure(gap_layers, territory_geom)
            _add_image_bytes(_fig_to_bytes(fig_gap), width_inches=6.0)
        except Exception as e:
            doc.add_paragraph(f'[Map unavailable: {e}]')

        doc.add_paragraph(
            'Red: strict gaps (no PA coverage). '
            'Amber: qualitative gaps (only weak protection classes). '
            'Blue: potential ecological corridors (within 5 km of two or more PA patches).'
        ).italic = False
        doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Section 4: Criterion Profiles
    # -----------------------------------------------------------------------
    if zonal_df is not None and len(zonal_df) > 0:
        doc.add_page_break()
        _heading('4. Criterion Profiles within Protected Areas', level=1)

        doc.add_paragraph(
            'Mean values of each MCE criterion raster within each IUCN category. '
            'All continuous criteria are min-max normalised to [0–1] for comparability. '
            'Anthropogenic pressure is inverted (low pressure = high score).'
        )

        _heading('4.1 Criterion Scores by IUCN Category', level=2)
        try:
            fig_crit = _criterion_bar_figure(zonal_df)
            _add_image_bytes(_fig_to_bytes(fig_crit), width_inches=6.5)
        except Exception as e:
            doc.add_paragraph(f'[Chart unavailable: {e}]')

        _heading('4.2 Summary Table', level=2)
        try:
            from modules.module1_protected_areas.zonal_stats import criterion_coverage_summary
            summary = criterion_coverage_summary(
                zonal_df[zonal_df['criterion'] != 'landuse']
            )
            summary_disp = summary.reset_index()
            for col in summary_disp.columns[1:]:
                summary_disp[col] = summary_disp[col].apply(
                    lambda x: f'{x:.3f}' if pd.notna(x) else '—'
                )
            _add_table(summary_disp)
        except Exception as e:
            doc.add_paragraph(f'[Table unavailable: {e}]')

        _heading('4.3 Detailed Statistics (min / median / max / std)', level=2)
        try:
            detailed = []
            for criterion in zonal_df['criterion'].unique():
                if criterion == 'landuse':
                    continue
                for _, row in zonal_df[zonal_df['criterion'] == criterion].iterrows():
                    detailed.append({
                        'Criterion':      criterion,
                        'IUCN Category':  row['iucn_cat'],
                        'Min':            f"{row['min']:.3f}",
                        'Median':         f"{row['median']:.3f}",
                        'Max':            f"{row['max']:.3f}",
                        'Std':            f"{row['std']:.3f}",
                        'Pixel Count':    f"{row['pixel_count']:,}",
                    })
            _add_table(pd.DataFrame(detailed))
        except Exception as e:
            doc.add_paragraph(f'[Table unavailable: {e}]')

    # -----------------------------------------------------------------------
    # Footer
    # -----------------------------------------------------------------------
    doc.add_page_break()
    footer = doc.add_paragraph(
        'OECM Favourability Tool — Module 1 Diagnostic Report\n'
        f'Generated {datetime.now().strftime("%Y-%m-%d %H:%M")} | '
        'Specification: CBD COP14 decision 14/8'
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    footer.runs[0].font.size = Pt(9)

    # -----------------------------------------------------------------------
    # Serialise to bytes
    # -----------------------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
